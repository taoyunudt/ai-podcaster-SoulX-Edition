# utils/document_analyzer.py - 文档和网页内容分析

import requests
from bs4 import BeautifulSoup
from docx import Document
import PyPDF2
from io import BytesIO
from typing import Optional
from utils.log_utils import info, error


class DocumentAnalyzer:
    """文档和网页内容分析器"""

    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'
        })

    def analyze_url(self, url: str) -> dict:
        """
        分析网页内容，提取主题

        Args:
            url: 网页地址

        Returns:
            dict: {title, content, theme}
        """
        try:
            info(f"🔍 正在分析网址: {url}")

            response = self.session.get(url, timeout=10)
            response.raise_for_status()

            soup = BeautifulSoup(response.text, 'html.parser')

            # 提取标题
            title = ''
            if soup.title:
                title = soup.title.string.strip()

            # 提取正文内容
            content = ''
            # 移除脚本和样式
            for script in soup(['script', 'style', 'nav', 'footer', 'header']):
                script.decompose()

            # 获取主要内容
            if soup.body:
                content = soup.body.get_text(separator='\n', strip=True)
                # 限制长度
                if len(content) > 5000:
                    content = content[:5000] + '...'

            info(f"✅ 网页分析完成: {title}")

            return {
                'title': title,
                'content': content,
                'url': url,
                'type': 'url'
            }

        except Exception as e:
            error(f"❌ 网页分析失败: {str(e)}")
            return None

    def analyze_word(self, file_path: str) -> dict:
        """
        分析 Word 文档

        Args:
            file_path: Word 文档路径

        Returns:
            dict: {title, content, theme}
        """
        try:
            info(f"🔍 正在分析 Word 文档: {file_path}")

            doc = Document(file_path)

            # 提取标题
            title = ''
            if doc.paragraphs:
                title = doc.paragraphs[0].text.strip()

            # 提取所有段落
            content = '\n'.join([para.text for para in doc.paragraphs])

            info(f"✅ Word 文档分析完成")

            return {
                'title': title or '未命名文档',
                'content': content,
                'file_path': file_path,
                'type': 'word'
            }

        except Exception as e:
            error(f"❌ Word 文档分析失败: {str(e)}")
            return None

    def analyze_pdf(self, file_path: str) -> dict:
        """
        分析 PDF 文档

        Args:
            file_path: PDF 文档路径

        Returns:
            dict: {title, content, theme}
        """
        try:
            info(f"🔍 正在分析 PDF 文档: {file_path}")

            content = ''
            title = ''

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # 提取标题（第一页的第一段）
                if pdf_reader.pages:
                    first_page = pdf_reader.pages[0]
                    text = first_page.extract_text()
                    if text:
                        lines = text.split('\n')
                        title = lines[0].strip()

                # 提取所有页面内容
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    content += text + '\n'

                # 限制长度
                if len(content) > 5000:
                    content = content[:5000] + '...'

            info(f"✅ PDF 文档分析完成")

            return {
                'title': title or '未命名文档',
                'content': content,
                'file_path': file_path,
                'type': 'pdf'
            }

        except Exception as e:
            error(f"❌ PDF 文档分析失败: {str(e)}")
            return None

    def extract_theme(self, content: str, model_name: str = "qwen-turbo") -> str:
        """
        使用大模型提取主题

        Args:
            content: 文档内容
            model_name: 模型名称

        Returns:
            str: 主题描述
        """
        try:
            from dashscope import Generation
            from config import DASHSCOPE_API_KEY

            prompt = f"""
请分析以下内容，提取出最适合制作播客的主题。

【要求】
1. 主题要具体、有趣
2. 适合两人对话的形式
3. 用一句话描述（不超过30字）

【内容】
{content[:2000]}

【主题】
"""

            response = Generation.call(
                model=model_name,
                prompt=prompt,
                api_key=DASHSCOPE_API_KEY,
                result_format='message'
            )

            if response.status_code == 200:
                if not response.output or not response.output.choices:
                    error("API返回结果格式错误")
                    return "通用主题"

                theme = response.output.choices[0].message.content.strip()
                info(f"✅ 主题提取成功: {theme}")
                return theme
            else:
                error(f"主题提取失败: {response.message}")
                return "通用主题"

        except Exception as e:
            error(f"❌ 主题提取失败: {str(e)}")
            return "通用主题"
